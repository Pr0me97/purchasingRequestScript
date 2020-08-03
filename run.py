import os

import xlrd
from docx import Document
from docx.oxml.ns import qn
from RmbBigMaker import RmbBigMaker

#准备写入内容

def getExcelRow(): #获取Excel表格数据的行数
    file='询价单价格汇总.xlsx'
    if os.path.exists(file):
        rd=xlrd.open_workbook(file)
        currentSheet=rd.sheet_by_name("价格汇总")
        row=currentSheet.nrows
        return row-1 #返回row行数,减去表头行
    else:
        print("询价单价格汇总表不存在！！！")

def readExcelData(row): #读取询价单价格汇总表的数据;传参row为行数,从1开始计数，0位表头
    file='询价单价格汇总.xlsx'
    if os.path.exists(file):
        rd=xlrd.open_workbook(file)
        currentSheet=rd.sheet_by_name("价格汇总")
        info=currentSheet.row_values(row)
        return info #返回row行的所有数据信息
    else:
        print("询价单价格汇总表不存在！！！")

def inputCompanyName(): #输入陪标公司名
    company1=input("请输入陪标公司1名称(回车确认)：")
    company2=input("请输入陪标公司2名称(回车确认)：")
    nameList=[company1,company2]
    return nameList

def editDocxWin(info): #修改中标模板；传参info为修改的信息，类型为list
    docx=Document('询价单-中标.docx')
    #读取并修改模板中的段落部分
    pars = docx.paragraphs
    docx.styles['Normal'].font.name = u'宋体'
    docx.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    
    if "XXXX" in pars[0].text:
        text=pars[0].text.replace("XXXX",info[0])
        pars[0].text=text
    
    #读取并修改模板中的表格部分
    table=docx.tables[0]

    cell=table.cell(1,1)
    p=cell.paragraphs[0]
    run=p.add_run(info[0])

    cell=table.cell(1,4)
    p=cell.paragraphs[0]
    run=p.add_run(str(info[1]/10000)+"万元")
    cell=table.cell(1,5)
    p=cell.paragraphs[0]
    run=p.add_run(str(info[1]/10000)+"万元")

    text=table.cell(2,1).text.replace("XXXX","%.2f"%info[1])
    text=text.replace("CCCC",RmbBigMaker(str(int(info[1]))))
    table.cell(2,1).text=text

    docx.save("询价单/%s_询价单_中标.docx"%info[0])

def editDocxCompany1(info,companyName): #修改中标模板；传参info为修改的信息，类型为list
    docx=Document('询价单-陪标1.docx')
    docx.styles['Normal'].font.name = u'黑体'
    docx.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    #读取并修改模板中的段落部分
    pars=docx.paragraphs
    text=pars[0].text.replace("XXXX",info[0])
    pars[0].text=text
    text=pars[2].text.replace("XXXX",info[0])
    pars[2].text=text
    text=pars[3].text.replace("XXXX",str(info[2]/10000))
    text=text.replace("CCCC",RmbBigMaker(str(int(info[2]))))
    pars[3].text=text
    text=pars[4].text.replace("XXXX",companyName)
    pars[4].text=text
    text=pars[13].text.replace("XXXX",companyName)
    pars[13].text=text
    docx.save("询价单/%s_%s_询价单.docx"%(info[0],companyName))

def editDocxCompany2(info,companyName): #修改中标模板；传参info为修改的信息，类型为list
    docx=Document('询价单-陪标2.docx')
    docx.styles['Normal'].font.name = u'楷体'
    docx.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
    pars=docx.paragraphs
    text=pars[0].text.replace("XXXX",info[0])
    pars[0].text=text
    text=pars[14].text.replace("XXXX",companyName)
    pars[14].text=text

    table=docx.tables[0]
    cell=table.cell(1,1)
    p=cell.paragraphs[0]
    run=p.add_run(info[0])       

    cell=table.cell(1,3)
    p=cell.paragraphs[0]
    run=p.add_run("%.2f元"%info[3])

    text=table.cell(2,1).text.replace("XXXX","%.2f"%info[3])
    text=text.replace("CCCC",RmbBigMaker(str(int(info[3]))))
    table.cell(2,1).text=text
    docx.save("询价单/%s_%s_询价单.docx"%(info[0],companyName))

def main():
    companyNameList=inputCompanyName() #输入2家陪标公司名
    rows=getExcelRow()
    flag=1
    while flag<=rows:
        info=readExcelData(flag)
        editDocxWin(info)
        editDocxCompany1(info,companyNameList[0])
        editDocxCompany2(info,companyNameList[1])
        flag+=1

main() 
#print(RmbBigMaker(str(input('请输输入数字：'))))